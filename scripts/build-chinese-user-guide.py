from __future__ import annotations

from datetime import date
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docs"
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "KSP_Mission_Display_中文使用教程_v0.1.2.docx"
SCREENSHOT = Path(
    os.environ.get("KMD_GUIDE_SCREENSHOT", ROOT / "docs" / "assets" / "display-reference.png")
)

FONT_CN = "Microsoft YaHei"
FONT_MONO = "IBM Plex Mono"
NAVY = "102638"
BLUE = "2E93C7"
CYAN = "65C7F2"
INK = "192C3A"
MUTED = "5D7180"
PALE = "E8F2F7"
PALE_BLUE = "DDEEF7"
PALE_YELLOW = "FFF4D6"
PALE_RED = "FCE8E8"
GREEN = "2F9E67"
YELLOW = "D69E2E"
RED = "C04A4A"
WHITE = "FFFFFF"
GRAY_LINE = "C9D7DF"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, *, name=FONT_CN, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRAY_LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_widths(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total_twips = int(sum(widths_in) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, (cell, width) in enumerate(zip(row.cells, widths_in)):
            width_twips = int(width * 1440)
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_twips))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_keep_with_next(p)
    return p


def add_body(doc, text, *, bold_prefix=None, color=INK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_step(doc, number, title, detail):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(title)
    set_run_font(r, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.375)
    p2.paragraph_format.space_after = Pt(7)
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(detail)
    set_run_font(r2, color=INK)


def add_callout(doc, label, text, *, kind="info"):
    colors = {
        "info": (PALE_BLUE, BLUE),
        "warning": (PALE_YELLOW, YELLOW),
        "danger": (PALE_RED, RED),
        "success": ("E6F5EC", GREEN),
    }
    fill, accent = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    set_table_borders(table, color=accent, size=10)
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_run_font(r, size=10, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(text)
    set_run_font(r2, size=10, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    return spacer


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.5])
    set_table_borders(table, color="8FA8B5", size=6)
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F3F6F8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name=FONT_MONO, size=8.6, color="20323D")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_fill(cell, PALE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=9.2, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(value))
            set_run_font(r, size=9.0, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (11.5, NAVY, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_CN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT_CN
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("KSP MISSION DISPLAY  /  中文使用教程")
    set_run_font(r, size=8.2, bold=True, color=MUTED)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), GRAY_LINE)
    pBdr.append(bottom)
    pPr.append(pBdr)

    footer = section.footer
    add_page_field(footer.paragraphs[0])


def build_document():
    doc = Document()
    setup_styles(doc)
    props = doc.core_properties
    props.title = "KSP Mission Display 中文使用教程"
    props.subject = "安装、kRPC 连接、任务创建、Display、FDO 与 Mission Planner"
    props.author = "KSP Mission Display Project"
    props.keywords = "KSP,RSS,RO,kRPC,FDO,GEO,TLI,遥测,任务规划"

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("KSP MISSION DISPLAY")
    set_run_font(r, name="Barlow Condensed", size=29, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(9)
    r = p2.add_run("中文使用教程")
    set_run_font(r, size=25, bold=True, color=BLUE)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(18)
    r = p3.add_run("RSS / RO · kRPC 实时遥测 · FDO · GEO Mission Planner")
    set_run_font(r, size=11.5, color=MUTED)

    if SCREENSHOT.exists():
        pimg = doc.add_paragraph()
        pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pimg.paragraph_format.space_after = Pt(8)
        pimg.add_run().add_picture(str(SCREENSHOT), width=Inches(6.45))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(22)
        r = cap.add_run("Display 主屏：任务名称、MET/UT、上面级状态、轨道参数与事件时间线")
        set_run_font(r, size=8.5, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    set_table_widths(meta, [1.45, 5.05])
    set_table_borders(meta, color=GRAY_LINE)
    for row, values in enumerate(
        [
            ("适用版本", "KSP Mission Display 0.1.2 / Prototype 0.2"),
            ("适用环境", "Windows x64 · KSP RSS/RO · kRPC"),
            ("文档日期", str(date.today())),
        ]
    ):
        for col, value in enumerate(values):
            cell = meta.rows[row].cells[col]
            if col == 0:
                set_cell_fill(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=9.2, bold=(col == 0), color=NAVY if col == 0 else INK)

    page_break(doc)

    # 1
    add_heading(doc, "1. 先读这一页：最快上手", 1)
    add_callout(
        doc,
        "一分钟结论",
        "安装版和便携版都已包含网页前端与 Python 后端，不需要另外安装 Python。要获得真实数据，必须先在 KSP 中启动 kRPC Server，并让程序连接到 50000/50001 端口。",
        kind="success",
    )
    add_heading(doc, "1.1 推荐启动顺序", 2)
    add_step(doc, 1, "启动 KSP 并进入飞行场景", "选择要监控的火箭或卫星作为当前 active vessel。")
    add_step(doc, 2, "开启 kRPC Server", "确认 RPC 端口为 50000、Stream 端口为 50001，并允许本机连接。")
    add_step(doc, 3, "启动 KSP Mission Display", "优先使用安装版；若不想安装，直接运行便携版 EXE。")
    add_step(doc, 4, "创建任务", "在 Mission Setup 填写任务名称、Profile、分级与事件定义，然后进入 Display/FDO。")
    add_step(doc, 5, "核对状态灯", "绿色 LIVE 表示 kRPC 实时数据；黄色 SIMULATION 表示没有实时信号或正在使用模拟回退。")

    add_heading(doc, "1.2 两种发行包", 2)
    add_table(
        doc,
        ["文件", "用途", "建议"],
        [
            ("KSP Mission Display Setup 0.1.2.exe", "Windows 安装程序，可选择目录并创建快捷方式", "日常使用"),
            ("KSP Mission Display 0.1.2.exe", "单文件便携版，不写入固定安装目录", "测试或移动使用"),
        ],
        [2.7, 2.55, 1.25],
    )
    add_callout(
        doc,
        "Windows SmartScreen",
        "当前版本没有代码签名，Windows 可能显示“未知发布者”。确认文件来自本项目 release 目录后，可选择“更多信息 -> 仍要运行”。",
        kind="warning",
    )

    # 2
    add_heading(doc, "2. kRPC 实时连接设置", 1)
    add_heading(doc, "2.1 数据链路", 2)
    add_table(
        doc,
        ["环节", "作用", "默认地址"],
        [
            ("KSP + kRPC", "提供飞行器、轨道、天体和事件数据", "127.0.0.1:50000 / 50001"),
            ("Python Backend", "读取 kRPC、计算派生量、输出 REST/WebSocket", "127.0.0.1:8021"),
            ("Electron/Web UI", "显示 Display、FDO、Mission Planner", "应用内本地网页"),
        ],
        [1.45, 3.45, 1.60],
    )
    add_body(doc, "实时遥测路径：KSP -> kRPC -> Python Backend -> 50 Hz WebSocket -> FDO/Display。")

    add_heading(doc, "2.2 KSP 端检查", 2)
    for item in (
        "kRPC 插件已正确安装，并能在游戏内打开 Server 窗口。",
        "RPC Port = 50000；Stream Port = 50001。",
        "本机使用时 Address 保持 127.0.0.1；局域网使用时允许对应网卡和防火墙规则。",
        "当前飞行器必须处于可控或可读取状态；程序默认绑定 active vessel。",
    ):
        add_bullet(doc, item)
    add_callout(
        doc,
        "只读安全边界",
        "当前版本不会控制节流、姿态、分级或机动节点。它只读取数据、记录事件并给出规划建议。",
        kind="info",
    )

    add_heading(doc, "2.3 状态灯含义", 2)
    add_table(
        doc,
        ["显示", "颜色", "含义", "处理"],
        [
            ("LIVE", "绿色", "kRPC 实时信号正常", "可使用 FDO 与实际 GEO Planner"),
            ("SIMULATION", "黄色", "没有实时信号或使用模拟回退", "检查 KSP、kRPC 和端口"),
            ("HOLD", "黄色/红色", "轨道条件不满足规划器要求", "先完成停车轨道或修正输入"),
        ],
        [1.1, 0.8, 2.25, 2.35],
    )

    page_break(doc)

    # 3
    add_heading(doc, "3. 创建任务与定义火箭结构", 1)
    add_heading(doc, "3.1 Mission Setup", 2)
    add_body(doc, "首次进入程序时先创建任务。任务配置决定 Display 的任务名称、飞行阶段、时间线以及 Mission Planner 的 Profile。")
    add_table(
        doc,
        ["字段", "填写内容"],
        [
            ("Mission Name", "任务名称，例如 H-IIA Flight 13、GEO Comms 01"),
            ("Mission Profile", "Earth Orbit、GEO Slot 或 TLI"),
            ("Vehicle / Flight", "火箭名称、批次或飞行编号"),
            ("Target", "目标轨道、目标槽位经度或月球任务约束"),
        ],
        [1.75, 4.75],
    )

    add_heading(doc, "3.2 Stage & Event Definitions", 2)
    add_body(doc, "分级和事件条目可以上下移动。程序会尝试从当前 craft 的发动机、activation stage、decouple stage、整流罩和资源数据推导结构，但推导结果应由用户复核。")
    add_bullet(doc, "助推器：点火、关机、分离。")
    add_bullet(doc, "一级：点火、MECO、级间分离。")
    add_bullet(doc, "二级/上面级：SEI、SECO、滑行、再次点火。")
    add_bullet(doc, "整流罩：抛罩事件。")
    add_bullet(doc, "载荷：星箭分离。")
    add_body(doc, "事件记录使用 MET，目标精度为 0.01 s；真实精度仍受 kRPC 采样率、游戏帧率和事件判据影响。")

    add_heading(doc, "3.3 自动推导的限制", 2)
    add_callout(
        doc,
        "必须人工检查",
        "RO 发动机点火次数、并联发动机、热分离、跨级供液、复杂整流罩和模组自定义模块可能无法仅靠 craft 数据可靠判断。自动推导是初稿，不是最终任务定义。",
        kind="warning",
    )

    # 4
    add_heading(doc, "4. Display 主屏", 1)
    add_body(doc, "Display 是第一屏 OSD，用于大屏、直播或任务总览。它强调任务身份、时间、当前飞行阶段和少量关键参数，不承担完整 FDO 分析。")
    add_table(
        doc,
        ["区域", "内容", "说明"],
        [
            ("顶部", "导航、LIVE/SIMULATION", "先确认数据源状态"),
            ("左侧", "任务名、飞行编号、当前阶段", "例如 Stage 2 Coast / TLI"),
            ("中央", "MET 与 UT 日期时间", "MET 大字；UT 为游戏世界时间换算显示"),
            ("右侧", "上面级状态", "推进剂、点火次数、剩余 Δv"),
            ("底部参数", "速度、高度、Ap、Pe、倾角、节点倒计时", "100 km 以上速度标签切换为 INERTIAL SPEED"),
            ("事件轴", "Liftoff、MAX-Q、MECO、SECO、Restart、Insertion", "实际事件显示实际 MET"),
        ],
        [1.2, 2.3, 3.0],
    )
    add_callout(
        doc,
        "MAX-Q",
        "进入最大动压阶段时，Display 应显示 MAX-Q 阶段名称以及最大动压数值。动压单位通常为 kPa，具体以当前界面标注为准。",
        kind="info",
    )

    page_break(doc)

    # 5
    add_heading(doc, "5. FDO 发射与轨道数据屏", 1)
    add_body(doc, "FDO 页面面向高时间分辨率的发射监视。实时模式使用 WebSocket Fast Channel，目标采样率为 50 Hz。")
    add_heading(doc, "5.1 关键数据", 2)
    add_table(
        doc,
        ["数据", "用途"],
        [
            ("Pitch / Heading / Roll", "监视姿态与制导趋势"),
            ("Surface / Inertial Speed", "低空使用地表速度；100 km 以上切换惯性速度"),
            ("Dynamic Pressure / Mach", "识别气动载荷和 MAX-Q"),
            ("Altitude / Ap / Pe", "监视上升轨迹与入轨质量"),
            ("Attitude Director", "显示姿态、目标方向和误差"),
            ("Sub-satellite Track", "绘制星下点连续轨迹；显示层使用平滑曲线"),
        ],
        [2.25, 4.25],
    )
    add_heading(doc, "5.2 链路质量指标", 2)
    add_bullet(doc, "Measured Rate：实际收到的帧率。")
    add_bullet(doc, "Gateway Latency：Python Gateway 到页面的延迟。")
    add_bullet(doc, "Sequence / Dropped Frames：序列号和丢帧数量。")
    add_bullet(doc, "FRAME CHECK：用于判断当前数据是否连续、可信。")
    add_callout(
        doc,
        "高频不等于高精度",
        "50 Hz 表示传输刷新率。传感量的物理精度仍受 KSP 仿真步长、游戏帧率、kRPC 数据源和模组影响。",
        kind="warning",
    )

    # 6
    add_heading(doc, "6. Mission Planner：GEO Slot", 1)
    add_body(doc, "GEO Planner 使用当前 active vessel 的实际停车轨道和天体常数，搜索未来 AN/DN 点火机会。它不会使用断线时的假候选。")
    add_heading(doc, "6.1 使用前提", 2)
    add_bullet(doc, "飞行器已经进入稳定停车轨道，Pe 为正且轨道数据有效。")
    add_bullet(doc, "kRPC 为绿色 LIVE。")
    add_bullet(doc, "当前 active vessel 就是要执行转移的上面级或组合体。")
    add_bullet(doc, "任务 Profile 选择 GEO Slot。")

    add_heading(doc, "6.2 人工输入", 2)
    add_table(
        doc,
        ["输入", "含义", "建议起始值"],
        [
            ("Target Longitude", "目标地固经度，东经 0-360°", "110.0°E"),
            ("Tolerance", "远地点星下点经度允许误差", "0.1°"),
            ("Node Filter", "搜索 AN、DN 或全部节点", "ALL"),
            ("Max Nodes", "向前搜索的节点数量", "12-30"),
        ],
        [1.65, 3.15, 1.70],
    )
    add_body(doc, "UT/MET、轨道根数、天体 μ、半径、角速度和同步半径均来自 kRPC，不需要手工抄写。")

    add_heading(doc, "6.3 结果怎么读", 2)
    add_table(
        doc,
        ["结果字段", "解释"],
        [
            ("AN / DN", "升交点或降交点点火机会"),
            ("UT Burn / MET Burn", "点火世界时与任务时间"),
            ("Injection Δv", "瞬时纯顺行模型下，抬远地点至同步半径所需 Δv"),
            ("Coast to Apogee", "点火后到远地点的滑行时间"),
            ("Apogee Longitude", "远地点时刻的 KSP 地固星下点经度"),
            ("Error / Window", "相对目标经度的误差与容许点火窗口"),
        ],
        [2.0, 4.5],
    )

    page_break(doc)

    # 7
    add_heading(doc, "7. GEO SLOT INSERTION 到底算什么", 1)
    add_body(doc, "本程序当前实现的是 L1_TWO_BODY_IMPULSIVE 建议模型。它回答的是：在未来哪个 AN/DN 附近进行理想瞬时顺行点火，可以让转移轨道远地点到达同步半径，并使远地点星下点尽量接近目标经度。")
    add_table(
        doc,
        ["计算步骤", "当前实现"],
        [
            ("节点时间", "由停车轨道 ArgPe/真近点角和 kRPC 的 UT 求得"),
            ("点火状态", "读取节点时刻的实际 KSP 非旋转参考系位置，并用邻近采样求速度"),
            ("注入 Δv", "对纯顺行瞬时 Δv 做二分求解，使远地点半径等于同步半径"),
            ("远地点传播", "两体轨道传播至新轨道远地点"),
            ("地固经度", "使用 KSP 当前参考系映射，不套现实 GMST/JD"),
            ("候选排序", "按目标经度误差绝对值排序"),
        ],
        [1.65, 4.85],
    )
    add_callout(
        doc,
        "不是自动变轨",
        "结果是 FDO 建议，不会创建或执行机动节点。有限燃烧、推力方向变化、联合平面变化、发动机点火延迟和质量变化尚未进入 L1 模型。",
        kind="danger",
    )

    add_heading(doc, "7.1 倾角问题", 2)
    add_body(doc, "如果 GTO 仍有明显倾角，即使远地点经度接近目标，也不代表已经完成真正的地球静止定点。最终 GEO 需要近圆轨道、同步周期并把倾角压到接近 0°。平面变化通常在远地点速度较低时执行。")
    add_heading(doc, "7.2 为什么会显示 HOLD", 2)
    add_bullet(doc, "KSP/kRPC 没有连接。")
    add_bullet(doc, "仍在发射台或亚轨道，Pe 为负。")
    add_bullet(doc, "轨道不是稳定椭圆，无法生成同步远地点候选。")
    add_bullet(doc, "当前 active vessel 不是预期的上面级。")

    # 8
    add_heading(doc, "8. TLI Profile 的当前边界", 1)
    add_callout(
        doc,
        "当前为 SIMULATED",
        "TLI 页面目前用于界面和数据结构演示，候选 C3、近月点和飞行时间尚未连接真实月球星历或 Principia 多体传播。不得把它当作实际点火解。",
        kind="warning",
    )
    add_body(doc, "后续真实 TLI 需要统一的地心惯性系、月球星历、多体/摄动传播、目标 B-plane 或近月点约束，以及有限燃烧模型。")

    # 9
    add_heading(doc, "9. 事件记录", 1)
    add_body(doc, "第二屏可记录实际发射事件。点火、起飞、助推器/一级/二级关机与分离、抛整流罩、再次点火和星箭分离均按 MET 记录。")
    add_table(
        doc,
        ["事件", "典型判据"],
        [
            ("LIFTOFF", "从 PRE_LAUNCH/ LANDED 转为飞行，并检测到速度或高度增长"),
            ("MECO / SECO", "对应发动机组推力由有效值降至阈值以下"),
            ("STAGE SEPARATION", "解耦器状态或 vessel/part 结构变化"),
            ("FAIRING JETTISON", "整流罩模块状态发生抛弃边沿"),
            ("PAYLOAD SEPARATION", "任务定义的目标解耦器或飞行器切换"),
        ],
        [2.0, 4.5],
    )
    add_callout(doc, "事件精度", "显示精度可以到 0.01 s，但事件判定精度不应高于原始采样和游戏物理步长。", kind="info")

    page_break(doc)

    # 10
    add_heading(doc, "10. 常见故障排查", 1)
    add_table(
        doc,
        ["现象", "可能原因", "解决方法"],
        [
            ("一直显示 SIMULATION", "kRPC 未开启或端口不可达", "检查 50000/50001、防火墙和 active vessel"),
            ("Planner 返回 503", "Python Backend 未连接 kRPC", "先启动 KSP/kRPC，再重新进入页面"),
            ("Planner 显示 HOLD", "未进入有效停车轨道", "确认 Pe>0、Ap/Pe 合理且轨道有效"),
            ("发射台 Pe 为负", "正常的未入轨轨道数据", "无需修复；入轨后再运行 GEO Planner"),
            ("FDO 帧率偏低", "游戏帧率低、CPU 忙或连接拥塞", "降低图形负载并检查 Dropped Frames"),
            ("网页能开但无实时数据", "8021 后端未启动或被占用", "检查后台日志和端口占用"),
            ("安装包被拦截", "开发版未签名", "核对文件来源和 SHA-256 后手动允许"),
        ],
        [1.65, 2.25, 2.60],
    )

    add_heading(doc, "10.1 日志", 2)
    add_body(doc, "桌面程序运行日志位于 Electron 应用日志目录，主要文件为：")
    add_bullet(doc, "kmd-web.log：内置网页服务器日志。")
    add_bullet(doc, "kmd-backend.log：Python/kRPC Gateway 日志。")
    add_body(doc, "发生启动错误时，先查看 kmd-backend.log 是否出现端口占用、模块缺失或 kRPC 连接拒绝。")

    add_heading(doc, "10.2 快速健康检查", 2)
    add_code(
        doc,
        "# 浏览器访问\nhttp://127.0.0.1:8021/health\n\n# 正常但未连接 KSP 时\nstatus: degraded\nkrpc_state: disconnected\n\n# 连接成功后\nkrpc_state: connected",
    )

    # 11
    add_heading(doc, "11. 开发者：手动构建", 1)
    add_heading(doc, "11.1 完整 Windows 构建", 2)
    add_code(doc, "cd ksp-mission-display\nnpm install\nnpm run electron:dist")
    add_body(doc, "完整构建依次执行：冻结 Python 后端 -> 构建独立网页 -> Electron Builder 生成安装版、便携版和 win-unpacked。")
    add_heading(doc, "11.2 常用命令", 2)
    add_table(
        doc,
        ["命令", "作用"],
        [
            ("npm run dev:lan", "在 0.0.0.0:3013 启动局域网开发服务器"),
            ("npm run backend:bundle", "只生成冻结 Python 后端"),
            ("npm run build", "只生成独立网页运行时"),
            ("npm run electron:dist", "生成完整 Windows 发行包"),
            ("npm run lint", "检查前端代码"),
        ],
        [2.45, 4.05],
    )
    add_callout(
        doc,
        "科学后端必须一起打包",
        "GEO Planner 使用 NumPy。构建脚本不得排除 numpy；否则安装包虽然能生成，kmd-backend.exe 会在启动时失败。",
        kind="danger",
    )

    page_break(doc)

    # 12
    add_heading(doc, "12. 参考系与数值解释", 1)
    add_body(doc, "KSP/kRPC 的天体参考系是左手坐标。当前程序将 KSP 原生数据转换到内部右手动力学系，并通过测试固定轴向映射。")
    add_table(
        doc,
        ["参考系", "用途"],
        [
            ("body.reference_frame", "随天体旋转的 KSP 地固系，用于当前地理量和界面真值"),
            ("body.non_rotating_reference_frame", "天体中心非旋转系，用于轨道状态与传播初值"),
            ("内部右手系", "科学计算、叉积、轨道要素和传播"),
        ],
        [2.45, 4.05],
    )
    add_callout(
        doc,
        "不要直接套现实 GMST",
        "RSS 的物理常数接近现实地球，但 KSP 的 UT、旋转零点和经度定义由游戏参考系决定。当前 GEO Planner 使用 kRPC 参考系映射，不直接把游戏 UT 当作现实 UTC/UT1。",
        kind="warning",
    )
    add_body(doc, "未来接入 Principia、多体引力或长时间传播时，必须明确惯性系 epoch、第三体星历、KSP 地固映射和时间尺度；不能只用 theta = theta0 + omega·dt 代替完整契约。")

    add_heading(doc, "13. 发射前检查清单", 1)
    checks = [
        "KSP 已进入飞行场景，active vessel 正确。",
        "kRPC Server 已开启，50000/50001 可访问。",
        "应用状态灯为绿色 LIVE。",
        "任务名称、Profile、分级和事件顺序已复核。",
        "Display 的 MET、UT、飞行阶段和 MAX-Q 数值合理。",
        "FDO Measured Rate 接近目标，Dropped Frames 无持续增长。",
        "使用 GEO Planner 前已进入有效停车轨道。",
        "已理解 L1 结果不包含有限燃烧和自动控制。",
        "TLI 页面仍为 SIMULATED，不用于真实点火。",
    ]
    for item in checks:
        add_bullet(doc, f"□ {item}")

    final_spacer = add_callout(doc, "版本定位", "0.1.2 为可运行原型：实时遥测与 GEO L1 建议可用；有限燃烧、多体传播及自动执行待后续开发。", kind="info")
    # Word requires a paragraph after a table. Keep the final one microscopic so
    # it does not create an otherwise empty trailing page.
    final_spacer.paragraph_format.space_before = Pt(0)
    final_spacer.paragraph_format.space_after = Pt(0)
    final_spacer.paragraph_format.line_spacing = Pt(1)
    final_run = final_spacer.add_run("")
    final_run.font.size = Pt(1)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
